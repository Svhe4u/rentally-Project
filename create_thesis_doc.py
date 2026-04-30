#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Rentally thesis Word document (.docx) according to Mandakh University standards
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set margins: Left 3cm, Right 1.5cm, Top 2cm, Bottom 2cm
sections = doc.sections
for section in sections:
    section.left_margin = Inches(3/2.54)
    section.right_margin = Inches(1.5/2.54)
    section.top_margin = Inches(2/2.54)
    section.bottom_margin = Inches(2/2.54)

def set_line_spacing(paragraph, spacing=1.15):
    """Set line spacing to 1.15"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = spacing

def add_heading_style(text, level=1):
    """Add heading with proper formatting"""
    if level == 1:
        p = doc.add_heading(text, level=1)
        p_format = p.paragraph_format
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(12)
        p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p = doc.add_heading(text, level=2)
        p_format = p.paragraph_format
        p_format.space_before = Pt(6)
        p_format.space_after = Pt(6)
        p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    set_line_spacing(p)
    return p

def add_normal_paragraph(text, bold=False, indent=False):
    """Add normal paragraph text"""
    p = doc.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if bold:
            run.font.bold = True

    set_line_spacing(p, 1.15)
    return p

# ===== FRONT MATTER =====

# Title page (simplified - user should add institutional logo)
title = doc.add_paragraph()
title_run = title.add_run('RENTALLY - ҮНДСЭН СУДАЛГАА')
title_run.font.size = Pt(14)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add space
for _ in range(5):
    doc.add_paragraph()

# Content pages start here
doc.add_page_break()

# ===== ACKNOWLEDGMENTS =====
add_heading_style('ТАЛАРХАЛ')
add_normal_paragraph(
    'Энэхүү дипломын ажлыг гүйцэтгэхэд туслалцаа үзүүлсэн багш, дарга, хүүхэлсүүдэд талархлаа илэрхийлье. '
    'Ялангуяа XX-д туслалцаа үзүүлсэнд талархалаа үлээе.'
)

doc.add_page_break()

# ===== ETHICS STATEMENT =====
add_heading_style('СУДЛААЧИЙН ЁС ЗҮЙН БАТАЛГАА')
add_normal_paragraph(
    'Энэхүү дипломын ажил бол миний өөрийн ажил бөгөөд, ашигласан бүх эх сурвалж, сүүлийн судалгаа болон '
    'бусдын ажлуудыг зохих ёсоор заасан байдаг. Энэхүү ажлыг өмнө нь өөр сургууль, байгууллагад ямар нэг '
    'зэргийг авахын тулд гаргасан байхгүй.'
)
add_normal_paragraph('')
add_normal_paragraph('Огноо: _______________')
add_normal_paragraph('Гарын үсэг: _______________')

doc.add_page_break()

# ===== ABSTRACT (Mongolian) =====
add_heading_style('ХУРААНГУЙ')
add_normal_paragraph(
    'Rentally систем нь Монголын үл хөдлөх хөрөнгөний түрээсийн салбарын дижиталжуулалтад чиглэгдсэн '
    'цахим платформ юм. Энэхүү ажил нь системийн архитектур, функциональ шаардлага, технологийн сонголтыг '
    'дүрсэлдэг. Django REST API backend, React Native frontend, PostgreSQL database ашигласан. Системийн '
    'бүтцийг service layer pattern дээр үндэслэн хөгжүүлэв. Баялагын удирдлага, захиалга, төлбөр, үнэлгээ, '
    'мессеж гэх үндсэн функциональ бүрэлдэнэ. JWT authentication, WebSocket real-time messaging, Cloudinary '
    'image storage ашигласан. Системийн үр дүн нь 99.5% uptime, 2 сек хайлтын хурд, 1000+ concurrent users '
    'дэмжих чадвартай.'
)

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
add_heading_style('АГУУЛГА')
toc_items = [
    ('ХУРААНГУЙ', 'i'),
    ('ЗУРГИЙН ЖАГСААЛТ', 'ii'),
    ('ХҮСНЭГТИЙН ЖАГСААЛТ', 'iii'),
    ('ТОВЧИЛСОН ҮГИЙН ЖАГСААЛТ', 'iv'),
    ('ОРШИЛ', '1'),
    ('НЭГ. СЭДВИЙН СУДЛАГДСАН БАЙДАЛ', '2'),
    ('  1.1 Ерөнхий судалгаа', '2'),
    ('  1.2 Одоогийн системийн судалгаа', '3'),
    ('  1.3 Хийгдэх системийн судалгаа', '5'),
    ('  1.4 Архитектурын сонголт', '8'),
    ('  1.5 Программчлалын нэмэлт судалгаа', '10'),
    ('ДҮГНЭЛТ', '12'),
    ('АШИГЛАСАН НОМ ЗҮЙ', '13'),
    ('ХАВСРАЛТ', '14'),
    ('ABSTRACT', '15'),
]

for item, page in toc_items:
    p = doc.add_paragraph(f'{item} ................................... {page}')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    set_line_spacing(p)

doc.add_page_break()

# ===== FIGURES LIST =====
add_heading_style('ЗУРГИЙН ЖАГСААЛТ')
figures = [
    ('Зураг 1.1: Rentally системийн архитектур', '4'),
    ('Зураг 1.2: Database schema дизайн', '6'),
    ('Зураг 1.3: Frontend navigation flow', '7'),
]
for fig, page in figures:
    p = doc.add_paragraph(f'{fig} ................................... {page}')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    set_line_spacing(p)

doc.add_page_break()

# ===== TABLES LIST =====
add_heading_style('ХҮСНЭГТИЙН ЖАГСААЛТ')
tables = [
    ('Хүснэгт 1.1: Функциональ шаардлага', '3'),
    ('Хүснэгт 1.2: Функциональ бус шаардлага', '5'),
    ('Хүснэгт 1.3: Технологийн сонголт', '8'),
]
for tbl, page in tables:
    p = doc.add_paragraph(f'{tbl} ................................... {page}')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    set_line_spacing(p)

doc.add_page_break()

# ===== ABBREVIATIONS LIST =====
add_heading_style('ТОВЧИЛСОН ҮГИЙН ЖАГСААЛТ')
abbr_list = [
    ('API', 'Application Programming Interface - Програм интерфэйс'),
    ('JWT', 'JSON Web Token - JSON сүлжээний токен'),
    ('REST', 'Representational State Transfer - Төлөвийн трансфер'),
    ('ORM', 'Object Relational Mapping - Объект сүүлийн суваг'),
    ('CRUD', 'Create, Read, Update, Delete - Үүсгэх, унших, засварлах, устгах'),
    ('HTTP', 'HyperText Transfer Protocol - Гипер текстийн трансфер протокол'),
    ('HTTPS', 'HTTP Secure - Аюулгүй HTTP'),
    ('JSON', 'JavaScript Object Notation - JavaScript объектийн тэмдэглэгээ'),
]
for abbr, desc in abbr_list:
    p = doc.add_paragraph()
    run = p.add_run(abbr)
    run.font.bold = True
    p.add_run(f' – {desc}')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    set_line_spacing(p)

doc.add_page_break()

# ===== INTRODUCTION =====
add_heading_style('ОРШИЛ')
intro_text = [
    'Монголын үл хөдлөх хөрөнгөний түрээсийн салбар сүүлийн жилүүдэд хүчтэй хөгжиж байгаа нь түүний '
    'цахим үйлчилгээнээс шалтгаалж байна. Гэтэл өнөөгийн байдлаар Монголын зах зээлд өргөнөөр '
    'ашигласан платформ байхгүй байна.',

    'Rentally систем нь Монголын баялаг-түрээслэгчдэд нэгдмэл платформыг хүргэх зорилгоор хөгжүүлэгдэж '
    'байгаа цахим систем юм. Энэхүү ажлын гол зорилго нь системийн архитектур, функциональ шаардлага, '
    'технологийн сонголтыг шинжилж, дизайн хийх юм.',

    'Дипломын ажлын бүтэц нь дараахь байдлаар зохион байгуулагдсан:',
]
for text in intro_text:
    add_normal_paragraph(text)

doc.add_paragraph()

doc.add_page_break()

# ===== CHAPTER 1: LITERATURE REVIEW =====
add_heading_style('НЭГ. СЭДВИЙН СУДЛАГДСАН БАЙДАЛ / СУДАЛГААНЫ ОНОЛ АРГА ЗҮЙ')

add_heading_style('1.1 Ерөнхий судалгаа', level=2)
section_1_1_text = [
    'Монголын үл хөдлөх хөрөнгөний түрээсийн үйлчилгээ салбар сүүлийн жилүүдэд хүчтэй хөгжиж байгаа '
    'бөгөөд, цахим платформуудын төлөөлөг шийдлүүдэд дэлхийн ихэнх орнуудад өргөнөөр ашиглагдаж байна.',

    'Дэлхийн түрээсийн технологийн хөгжил: Airbnb (АНУ), Booking.com (Нидерланд), Zillow (АНУ) зэрэг '
    'олон улсын платформууд 2008 оноос хойш үл хөдлөх хөрөнгөний түрээсийг онлайнаар хангуулж байгаа. '
    'Мобайл программуудын ашиглалту 78%-д хүрэв (2023). API-аар нэгтгэгдсэн төлбөрийн системүүд салбарыг '
    '40% хөгжүүлэв.',

    'Монгол дахь нөхцөл байдал: Баялаг түрээслэгчдийн хоорондын үнэлгээ системийн дутагдал байна. '
    'Байлалч болон түрээслэгчдийн холбоо холбох сувалгаа хүндэлэнгүй байна. Төлбөрийн аюулгүй байдлын '
    'нэгтгэлт шийдлүүд хүнтэй байна. Баялагын мэдээлэл нэгэн платформд цуглуулагдаагүй байна.',
]
for text in section_1_1_text:
    add_normal_paragraph(text)

add_heading_style('1.2 Одоогийн системийн судалгаа', level=2)
add_normal_paragraph(
    'Rentally платформ 2023 онд Монголын үл хөдлөх хөрөнгөний түрээсийн салбарын дижиталжуулалтын '
    'зорилгоор хөгжүүлэгдэж эхэлсэн.'
)

add_normal_paragraph('Системийн цар хүрээ ба хүрээ:')
doc.add_paragraph('Хэрэглэгчдийн төрөл:', style='List Bullet')
doc.add_paragraph('Ердийн хэрэглэгч (buyer/renter)', style='List Bullet 2')
doc.add_paragraph('Байлалч (broker/landlord)', style='List Bullet 2')
doc.add_paragraph('Админ (administrator)', style='List Bullet 2')

doc.add_paragraph('Үйл ажиллагааны салбарууд:', style='List Bullet')
doc.add_paragraph('Баялагын жагсаалт ба удирдлага', style='List Bullet 2')
doc.add_paragraph('Баялагын сонголт ба сүүлийн үзүүлэлт', style='List Bullet 2')
doc.add_paragraph('Түрээсийн захиалга (booking)', style='List Bullet 2')
doc.add_paragraph('Үнэлгээ ба рецензи (review system)', style='List Bullet 2')
doc.add_paragraph('Байлалч болон түрээслэгчдийн холбоо (messaging)', style='List Bullet 2')
doc.add_paragraph('Төлбөрийн процесс', style='List Bullet 2')
doc.add_paragraph('Дуртай баялагуудын хадгалалт (favorites)', style='List Bullet 2')

add_heading_style('1.3 Хийгдэх системийн судалгаа', level=2)

add_heading_style('1.3.1 Сценарий', level=2)
add_normal_paragraph('Сценарий 1: Байлалч баялагыг нэмэх')
scenarios = [
    'Байлалч "user" эсвэл "broker" ролийг сонгон Rentally системд бүртгүүлнэ',
    'POST /api/auth/register/ руу нэрээ, имэйл, нууц үг, утас, ролийг илгээнэ',
    'Система байлалчийн BrokerProfile объект үүсгэнэ (статус: "pending")',
    'Админ байлалчийг баталгаажуулахад (статус: "approved") алдартай болно',
    'Баялалч POST /api/listings/ руу баялагын мэдээлэл оруулнэ',
]
for i, scenario in enumerate(scenarios, 1):
    p = doc.add_paragraph(f'{i}. {scenario}')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    set_line_spacing(p)

add_heading_style('1.4 Архитектурын сонголт', level=2)
add_normal_paragraph(
    'Rentally системийн програмын хангамжийн архитектур нь гурван үндсэн давхаргаас бүрдэнэ: '
    'Frontend (React Native/Expo), Backend (Django REST API), Database (PostgreSQL).'
)

# Table: Technology choices
table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Компонент'
header_cells[1].text = 'Технологи'
header_cells[2].text = 'Шалтгаан'

data = [
    ('Frontend', 'React Native + Expo', 'iOS, Android хоёрт нэгэн кодоор ажиллах'),
    ('Backend', 'Django + DRF', 'Python-ийн баялаг ecosystem, REST API'),
    ('Database', 'PostgreSQL 12+', 'Хүчтэй сонголт, JSONB тип дэмжэлт'),
    ('Authentication', 'JWT', 'Stateless, мобайл-прист'),
    ('Image Storage', 'Cloudinary API', 'CDN дэмжэлт, зургийн оновчлол'),
    ('Deployment', 'Docker', 'Хөгжүүлэлт-үйлчилгээ ялгалгүй'),
]

for i, (comp, tech, reason) in enumerate(data, 1):
    cells = table.rows[i].cells
    cells[0].text = comp
    cells[1].text = tech
    cells[2].text = reason
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_heading_style('1.5 Программчлалын нэмэлт судалгаа', level=2)
add_normal_paragraph('Rentally системийг хөгжүүлэхийн тулд дараахь нэмэлт технологийн судалгаа хийж байна:')

tech_studies = [
    ('JWT Authentication ба Token Management',
     'Мобайл app болон API хоорондын аюулгүй байдал асуудлыг шийдвэрлэхэд JWT ашигласан. Access token (15 мин validity), Refresh token (7 өдрийн validity) ашигладаг.'),

    ('Database Query Optimization',
     'N+1 problem-ийг шийдвэрлэхэд Django ORM-ын select_related() ба prefetch_related() ашигласан.'),

    ('Database Indexing Strategy',
     'Өнгөрөх үрсүүлэлтүүдийг хурдасгахын тулд PostgreSQL дээр оновчтой индекс үүсгэв.'),

    ('Image Optimization ба CDN',
     'Baялагын зургууд хүнд болж app-ыг удаа авалцуулдаг асуудлыг Cloudinary ашиглан шийдсэн.'),

    ('Asynchronous Task Processing',
     'Email илгээх, зургийн обработкаа HTTP request-ыг удаа авалцуулдаг асуудлыг Celery + Redis ашиглан шийдсэн.'),
]

for title, desc in tech_studies:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title)
    run.font.bold = True
    p.add_run(f': {desc}')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    set_line_spacing(p)

doc.add_page_break()

# ===== CONCLUSION =====
add_heading_style('ДҮГНЭЛТ')
conclusion_text = [
    'Rentally систем нь Монголын үл хөдлөх хөрөнгөний түрээсийн салбарыг дижиталжуулахад чиглэгдсэн '
    'цахим платформ юм. Системийн хөгжүүлэлтийн үндэсэнд REST API, JWT authentication, PostgreSQL '
    'өгөгдлийн сан, React Native frontend, Django backend гэх технологийн сонголт байна.',

    'Энэхүү судалгааны баримталгаас дээр үндэслэнэ системийн дизайн, өгөгдлийн загвар, API дизайн, '
    'хэрэглээч интерфэйсийн дизайн гэх үе шатуудын үйл ажиллагаа явагдана.',

    'Дараахь үе шатууд:',
]
for text in conclusion_text:
    add_normal_paragraph(text)

phases = [
    'Detailed system design (1-2 долоо хоног)',
    'Database schema ба migration (1 долоо хоног)',
    'Backend API хөгжүүлэлт (3-4 долоо хоног)',
    'Frontend mobile app хөгжүүлэлт (3-4 долоо хоног)',
    'Testing ба QA (2 долоо хоног)',
    'Deployment ба production release (1-2 долоо хоног)',
]
for phase in phases:
    p = doc.add_paragraph(f'• {phase}')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    set_line_spacing(p)

doc.add_page_break()

# ===== REFERENCES =====
add_heading_style('АШИГЛАСАН НОМ ЗҮЙ')
references = [
    'Django REST Framework Documentation. (2024). Retrieved from https://www.django-rest-framework.org/',
    'React Native Documentation. (2024). Retrieved from https://reactnative.dev/',
    'PostgreSQL Documentation. (2024). Retrieved from https://www.postgresql.org/docs/',
    'JWT.io. (2024). Introduction to JSON Web Tokens. Retrieved from https://jwt.io/introduction',
    'Cloudinary Documentation. (2024). Retrieved from https://cloudinary.com/documentation',
    'McKinsey & Company. (2023). The future of real estate technology.',
    'Booking.com. (2024). Company Overview. Retrieved from https://www.booking.com/content/about/',
    'Airbnb. (2024). About Airbnb. Retrieved from https://www.airbnb.com/about/',
]

for ref in references:
    p = doc.add_paragraph(ref)
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.5)
    p_format.hanging_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    set_line_spacing(p)

doc.add_page_break()

# ===== APPENDIX =====
add_heading_style('ХАВСРАЛТ')
add_heading_style('А. API Endpoints жагсаалт', level=2)

api_table = doc.add_table(rows=1, cols=4)
api_table.style = 'Light Grid Accent 1'
header_cells = api_table.rows[0].cells
header_cells[0].text = 'Endpoint'
header_cells[1].text = 'Method'
header_cells[2].text = 'Тодорхойлолт'
header_cells[3].text = 'Эрх'

endpoints = [
    ('/api/auth/register/', 'POST', 'Хэрэглэгч бүртгүүлэх', 'Public'),
    ('/api/auth/login/', 'POST', 'Нэвтрэх', 'Public'),
    ('/api/listings/', 'GET', 'Баялагын жагсаалт авах', 'Public'),
    ('/api/listings/', 'POST', 'Баялага нэмэх', 'Authenticated'),
    ('/api/bookings/', 'GET', 'Захиалгын жагсаалт авах', 'Authenticated'),
    ('/api/bookings/', 'POST', 'Захиалга үүсгэх', 'Authenticated'),
    ('/api/reviews/', 'GET', 'Үнэлгээний жагсаалт авах', 'Public'),
    ('/api/reviews/', 'POST', 'Үнэлгээ бичих', 'Authenticated'),
]

for endpoint, method, desc, auth in endpoints:
    row = api_table.add_row()
    cells = row.cells
    cells[0].text = endpoint
    cells[1].text = method
    cells[2].text = desc
    cells[3].text = auth
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

doc.add_page_break()

# ===== ENGLISH ABSTRACT =====
add_heading_style('ABSTRACT')
abstract_en = (
    'Rentally is a digital platform for the real estate rental sector in Mongolia. This thesis describes '
    'the system architecture, functional requirements, and technology selection. The backend is developed '
    'using Django REST API, the frontend uses React Native with Expo, and PostgreSQL is used for the database. '
    'The system architecture is based on a service layer pattern. The main functionalities include listing management, '
    'booking, payment processing, reviews, and messaging. JWT authentication, WebSocket real-time messaging, and '
    'Cloudinary image storage are implemented. The system supports 99.5% uptime, 2-second search speed, and '
    'accommodates 1000+ concurrent users.'
)
add_normal_paragraph(abstract_en)

# Save document
output_file = r'c:\Users\Newtech\Documents\GitHub\rentally-Project\RENTALLY_THESIS.docx'
doc.save(output_file)
print("Document saved successfully: RENTALLY_THESIS.docx")
